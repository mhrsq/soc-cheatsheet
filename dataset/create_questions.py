#!/usr/bin/env python3
"""Generate SOC Cyber Drill — 20 Questions & Answers DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    doc.styles[f'Heading {level}'].font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

# ─── TITLE ──────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SOC Cyber Drill Exercise')
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)
r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('20 Investigation Questions')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.add_run('PT Nusantara Digital — APT Attack Investigation\n').font.size = Pt(11)
t3.add_run('JagoanSiber · SOC Analyst Training\n').font.size = Pt(11)
t3.add_run('April 2026').font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Splunk URL: ').bold = True
p.add_run('http://141.94.36.151:8001')
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Login: ').bold = True
p2.add_run('admin / JagoanSiber2026')

doc.add_page_break()

# ─── INSTRUCTIONS ───────────────────────────────────────
doc.add_heading('Instructions', level=1)
doc.add_paragraph(
    'Kamu adalah SOC Analyst di PT Nusantara Digital. Tim keamanan mendeteksi beberapa anomali di SIEM '
    'yang mengindikasikan adanya serangan. Tugasmu adalah menginvestigasi menggunakan Splunk dan menjawab '
    '20 pertanyaan berikut secara berurutan. Pertanyaan disusun dari mudah ke sulit — setiap jawaban akan '
    'membantu kamu memahami gambaran besar insiden ini.'
)

doc.add_heading('Available Indexes', level=2)
indexes = [
    ('firewall', 'pan_traffic', 'Palo Alto firewall logs'),
    ('dns', 'pan_dns', 'DNS query logs'),
    ('ids', 'suricata', 'Suricata IDS alerts'),
    ('windows', 'WinEventLog', 'Windows Security Event Log'),
    ('sysmon', 'XmlWinEventLog', 'Sysmon events'),
    ('linux', 'linux_audit, syslog', 'Linux auditd + syslog'),
    ('web', 'access_combined', 'Apache access logs'),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 1'
h = tbl.rows[0].cells
h[0].text = 'Index'; h[1].text = 'Sourcetype'; h[2].text = 'Description'
for idx, st, desc in indexes:
    r = tbl.add_row().cells
    r[0].text = idx; r[1].text = st; r[2].text = desc

doc.add_paragraph()
doc.add_paragraph('Time range: 7-9 April 2026. Set Splunk time picker ke "All time" atau "Since April 7, 2026".')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Tips: ').bold = True
p.add_run('Gunakan operator | stats, | table, | timechart, | where, | sort untuk menganalisis hasil.')

doc.add_page_break()

# ─── QUESTIONS (Student Version) ────────────────────────
doc.add_heading('Questions', level=1)

questions = [
    # EASY (1-7) — Reconnaissance & Basic Search
    {
        'num': 1, 'level': '🟢 Easy', 'category': 'Reconnaissance',
        'q': 'Di index "ids" (Suricata), ada alert bertipe "ET SCAN". Dari IP address mana serangan scanning ini berasal?',
        'hint': 'Coba search: index=ids "ET SCAN"',
        'answer': '91.215.85.142',
        'query': 'index=ids "ET SCAN" | stats count by src_ip',
        'explanation': 'IP 91.215.85.142 (GeoIP: Romania) melakukan port scanning terhadap web server. Semua traffic dari IP ini diblok oleh firewall.',
    },
    {
        'num': 2, 'level': '🟢 Easy', 'category': 'Reconnaissance',
        'q': 'Berapa banyak port yang di-scan oleh attacker tersebut? Sebutkan minimal 5 port yang ditargetkan.',
        'hint': 'Coba filter berdasarkan IP attacker dan lihat dest_port',
        'answer': '18 port. Termasuk: 21, 22, 23, 25, 80, 443, 445, 3306, 3389, 8080',
        'query': 'index=ids src_ip=91.215.85.142 | stats count by dest_port | sort dest_port',
        'explanation': 'Attacker melakukan full port scan terhadap target. Ini adalah teknik reconnaissance (MITRE T1595).',
    },
    {
        'num': 3, 'level': '🟢 Easy', 'category': 'Reconnaissance',
        'q': 'Selain port scan, attacker juga melakukan web vulnerability scanning. Jenis serangan web apa yang dicoba? (Lihat di index "ids" dan "web")',
        'hint': 'Search: index=ids src_ip=91.215.85.142 "SQL" OR "Cross-Site" OR "Traversal"',
        'answer': 'SQL Injection (SELECT FROM, UNION SELECT), Cross-Site Scripting (XSS), Directory Traversal',
        'query': 'index=ids src_ip=91.215.85.142 event_type=alert | stats count by alert.signature',
        'explanation': 'Attacker menggunakan automated scanner (Nikto) untuk mencari vulnerability di web server.',
    },
    {
        'num': 4, 'level': '🟢 Easy', 'category': 'Reconnaissance',
        'q': 'Apakah web scanning tersebut berhasil? Cek di Apache access log — status code apa yang diterima attacker?',
        'hint': 'Search: index=web 91.215.85.142',
        'answer': 'Tidak berhasil. Status code: 404 (Not Found) dan 400 (Bad Request)',
        'query': 'index=web 91.215.85.142 | stats count by status',
        'explanation': 'Scanner gagal menemukan vulnerability yang bisa dieksploitasi. Attacker kemudian beralih ke vektor serangan lain: phishing.',
    },
    {
        'num': 5, 'level': '🟢 Easy', 'category': 'Initial Access',
        'q': 'Attacker beralih ke phishing. Di Sysmon (index=sysmon), cari proses yang di-spawn oleh Microsoft Word (WINWORD.EXE). Host mana yang terinfeksi dan siapa user-nya?',
        'hint': 'Search: index=sysmon EventID=1 ParentImage="*WINWORD*"',
        'answer': 'Host: WS-FINANCE-01.nusantara.local, User: NUSANTARA\\ratna.finance',
        'query': 'index=sysmon EventID=1 ParentImage="*WINWORD*" | table _time Computer Image CommandLine User ParentImage',
        'explanation': 'User ratna.finance membuka dokumen malicious (Invoice_Q1_2026.docm) yang berisi macro. Macro menjalankan cmd.exe → powershell.exe.',
    },
    {
        'num': 6, 'level': '🟢 Easy', 'category': 'Initial Access',
        'q': 'Apa nama file dokumen yang dibuka oleh user tersebut? (Lihat di ParentCommandLine dari proses WINWORD)',
        'hint': 'Cari di field ParentCommandLine atau CommandLine yang mengandung ".doc"',
        'answer': 'Invoice_Q1_2026.docm',
        'query': 'index=sysmon EventID=1 ParentImage="*WINWORD*" | table ParentCommandLine',
        'explanation': 'File .docm (macro-enabled document) adalah format yang sering dipakai untuk phishing karena bisa menjalankan macro VBA.',
    },
    {
        'num': 7, 'level': '🟢 Easy', 'category': 'Execution',
        'q': 'PowerShell yang dijalankan menggunakan parameter -enc (encoded). Cari command PowerShell yang encoded tersebut. Apa yang dilakukan script ini? (Hint: decode Base64-nya)',
        'hint': 'Search: index=sysmon EventID=1 Image="*powershell*" CommandLine="*enc*"',
        'answer': 'Download payload dari https://dl.techsupport-cdn.net/update.exe menggunakan IEX (New-Object Net.WebClient).DownloadString()',
        'query': 'index=sysmon EventID=1 Image="*powershell*" CommandLine="*enc*" | table _time CommandLine',
        'explanation': 'Base64 decode menunjukkan PowerShell downloader — teknik umum untuk menghindari deteksi (MITRE T1059.001, T1027).',
    },

    # MEDIUM (8-14) — C2, Lateral Movement, Credential Access
    {
        'num': 8, 'level': '🟡 Medium', 'category': 'C2 Communication',
        'q': 'Payload yang didownload membuat koneksi keluar. Di Sysmon Event ID 3 (network connection), ke IP dan port berapa payload tersebut berkomunikasi?',
        'hint': 'Search: index=sysmon EventID=3 Image="*svchost_update*"',
        'answer': 'IP: 45.77.65.211, Port: 443 (HTTPS)',
        'query': 'index=sysmon EventID=3 Image="*svchost_update*" | table _time DestinationIp DestinationPort Image',
        'explanation': 'Ini adalah C2 (Command & Control) server. Attacker menggunakan port 443 (HTTPS) supaya blend in dengan traffic normal.',
    },
    {
        'num': 9, 'level': '🟡 Medium', 'category': 'C2 Communication',
        'q': 'Domain apa yang digunakan untuk C2? Cari di DNS logs.',
        'hint': 'Search: index=dns 45.77.65.211',
        'answer': 'update-service.cloud',
        'query': 'index=dns 45.77.65.211 | stats count by domain',
        'explanation': 'C2 domain dibuat mirip dengan update service yang legitimate — teknik social engineering pada domain name.',
    },
    {
        'num': 10, 'level': '🟡 Medium', 'category': 'C2 Communication',
        'q': 'Berapa total koneksi C2 beacon yang terjadi selama 3 hari? Gunakan firewall log. Apakah ada pola waktu yang terlihat (beaconing)?',
        'hint': 'Search: index=firewall dest_ip=45.77.65.211 | timechart count span=5m',
        'answer': '~3,071 koneksi. Pola: regular interval setiap ~60 detik (beaconing pattern)',
        'query': 'index=firewall 45.77.65.211 | stats count\nindex=firewall 45.77.65.211 | timechart count span=5m',
        'explanation': 'Beaconing pattern = koneksi keluar ke IP yang sama dengan interval reguler. Ini signature khas C2 communication.',
    },
    {
        'num': 11, 'level': '🟡 Medium', 'category': 'Credential Access',
        'q': 'Sysmon Event ID 10 mendeteksi akses ke proses sensitif. Proses apa yang diakses dan ini indikasi teknik serangan apa?',
        'hint': 'Search: index=sysmon EventID=10',
        'answer': 'TargetImage: lsass.exe. Ini indikasi credential dumping menggunakan Mimikatz (MITRE T1003.001)',
        'query': 'index=sysmon EventID=10 | table _time Computer SourceImage TargetImage GrantedAccess',
        'explanation': 'LSASS (Local Security Authority Subsystem Service) menyimpan credentials di memory. Akses ke LSASS = attacker mencuri password/hash.',
    },
    {
        'num': 12, 'level': '🟡 Medium', 'category': 'Lateral Movement',
        'q': 'Setelah mencuri credentials, attacker login ke server lain. Di Windows Security Log, cari logon type 3 (network logon) yang berasal dari WS-FINANCE-01 (10.10.2.50) ke DC-01. Authentication package apa yang dipakai?',
        'hint': 'Search: index=windows EventCode=4624 10.10.2.50 DC-01',
        'answer': 'Authentication Package: NTLM (Pass-the-Hash attack)',
        'query': 'index=windows EventCode=4624 10.10.2.50 host="DC-01*" | table _time src_ip Account_Name LogonType',
        'explanation': 'NTLM authentication dari workstation ke DC = suspicious. Seharusnya pakai Kerberos di domain environment. NTLM mengindikasikan Pass-the-Hash.',
    },
    {
        'num': 13, 'level': '🟡 Medium', 'category': 'Lateral Movement',
        'q': 'Attacker juga bergerak ke Linux server. Cari di syslog: dari IP mana ada SSH login ke DB-SERVER-01 yang mencurigakan?',
        'hint': 'Search: index=linux sourcetype=syslog host=DB-SERVER-01 "Accepted"',
        'answer': 'SSH login dari 10.10.2.50 (WS-FINANCE-01) ke DB-SERVER-01 sebagai root menggunakan password',
        'query': 'index=linux sourcetype=syslog "Accepted" host=DB-SERVER-01 | table _time host',
        'explanation': 'Login SSH dari Windows workstation ke Linux DB server = tidak biasa. Apalagi sebagai root dengan password (bukan key). Ini lateral movement via SSH (T1021.004).',
    },
    {
        'num': 14, 'level': '🟡 Medium', 'category': 'Discovery',
        'q': 'Sebelum lateral movement, attacker menjalankan beberapa perintah discovery di WS-FINANCE-01. Sebutkan minimal 3 perintah yang dijalankan.',
        'hint': 'Search: index=sysmon EventID=1 Computer="WS-FINANCE-01*" (Image="*whoami*" OR Image="*ipconfig*" OR Image="*net.exe*" OR Image="*nltest*" OR Image="*systeminfo*")',
        'answer': 'whoami /all, ipconfig /all, net group "Domain Admins" /domain, nltest /domain_trusts, systeminfo',
        'query': 'index=sysmon EventID=1 Computer="WS-FINANCE-01*" Image="*whoami*" OR Image="*ipconfig*" OR Image="*net.exe*" OR Image="*nltest*" OR Image="*systeminfo*" | table _time Image CommandLine',
        'explanation': 'Ini adalah fase Discovery (MITRE T1082, T1069, T1482) — attacker mengumpulkan informasi tentang environment sebelum bergerak lateral.',
    },

    # HARD (15-20) — Persistence, Exfiltration, Defense Evasion, Timeline
    {
        'num': 15, 'level': '🔴 Hard', 'category': 'Persistence',
        'q': 'Attacker membuat akun backdoor di Domain Controller. Apa nama akun tersebut dan group apa yang ditambahkan?',
        'hint': 'Search: index=windows EventCode=4720 (user created) dan EventCode=4728 (added to group)',
        'answer': 'Account: svc_backup, ditambahkan ke group Domain Admins',
        'query': 'index=windows EventCode=4720 | table _time host Account_Name Target_Account_Name\nindex=windows EventCode=4728 | table _time Group_Name Member_Name',
        'explanation': 'Attacker membuat akun backdoor dengan nama yang terlihat legitimate (svc_backup) dan memberikan Domain Admin privileges — persistence technique (T1136.002).',
    },
    {
        'num': 16, 'level': '🔴 Hard', 'category': 'Persistence',
        'q': 'Selain backdoor account, attacker juga membuat scheduled task untuk persistence. Apa nama task-nya dan binary apa yang dijalankan?',
        'hint': 'Search: index=windows EventCode=4698 dan index=sysmon Image="*schtasks*"',
        'answer': 'Task: \\Microsoft\\Windows\\Maintenance\\SystemUpdate, menjalankan svchost_update.exe (malware) sebagai SYSTEM setiap hari jam 09:00',
        'query': 'index=sysmon EventID=1 Image="*schtasks*" | table _time CommandLine',
        'explanation': 'Scheduled task memberikan persistence — malware akan tetap jalan meskipun user logout. Nama task dibuat mirip Windows update supaya tidak dicurigai (T1053.005).',
    },
    {
        'num': 17, 'level': '🔴 Hard', 'category': 'Collection & Exfiltration',
        'q': 'Data apa yang dicuri dari DB-SERVER-01? Cari perintah yang dijalankan di auditd logs. Jelaskan langkah-langkah exfiltration yang dilakukan.',
        'hint': 'Search: index=linux sourcetype=linux_audit comm="mysqldump" OR comm="tar" OR comm="curl"',
        'answer': 'Langkah: 1) mysqldump --all-databases (dump semua database), 2) tar czf /tmp/bak.tar.gz (compress), 3) curl POST upload ke 185.234.216.88 (exfiltrate)',
        'query': 'index=linux sourcetype=linux_audit (comm="mysqldump" OR comm="tar" OR comm="curl") | table _time host comm exe\nindex=linux sourcetype=linux_audit a0="curl" | table _time',
        'explanation': 'Attacker melakukan data exfiltration dalam 3 langkah: dump → compress → upload. Transfer dilakukan jam 02:30 (after hours) untuk menghindari deteksi.',
    },
    {
        'num': 18, 'level': '🔴 Hard', 'category': 'Exfiltration',
        'q': 'Berapa besar data yang diexfiltrate? Cari di firewall log transfer ke IP staging attacker (185.234.216.88).',
        'hint': 'Search: index=firewall 185.234.216.88 | table bytes_sent',
        'answer': '~250 MB (262,000,000 bytes) — transfer besar di jam 02:30 WIB',
        'query': 'index=firewall dest_ip=185.234.216.88 | table _time src_ip dest_ip bytes_sent bytes_received',
        'explanation': 'Transfer 250MB ke IP di Russia pada jam 2 pagi = strong indicator of data exfiltration. Ini seharusnya trigger alert di SOC.',
    },
    {
        'num': 19, 'level': '🔴 Hard', 'category': 'Defense Evasion',
        'q': 'Attacker mencoba menutupi jejaknya. Dua teknik anti-forensics apa yang dilakukan? Sebutkan event log evidence-nya.',
        'hint': 'Search: index=windows EventCode=1102 dan index=linux comm="rm" dan index=sysmon Image="*wevtutil*"',
        'answer': '1) Cleared Windows Security event log (EventCode 1102, wevtutil cl Security pada 02:40). 2) Deleted staged file /tmp/bak.tar.gz (rm -f pada 02:35)',
        'query': 'index=windows EventCode=1102 | table _time host Account_Name\nindex=linux sourcetype=linux_audit comm="rm" | table _time\nindex=sysmon Image="*wevtutil*" | table _time CommandLine',
        'explanation': 'Clearing event log (T1070.001) dan menghapus file (T1070.004) adalah teknik defense evasion. Ironisnya, event "log cleared" (1102) itu sendiri tercatat — karena itu event terakhir sebelum log di-clear.',
    },
    {
        'num': 20, 'level': '🔴 Hard', 'category': 'Full Timeline',
        'q': 'FINAL: Susun timeline lengkap serangan ini dari awal sampai akhir. Sebutkan: waktu, teknik, host yang terlibat, dan MITRE ATT&CK ID untuk setiap tahap.',
        'hint': 'Gunakan semua jawaban sebelumnya untuk menyusun kronologis',
        'answer': (
            'Timeline:\n'
            '1. Apr 7, 09:15 — Reconnaissance: Port scan + web scan dari 91.215.85.142 → WEB-SERVER-01 [T1595]\n'
            '2. Apr 7, 14:47 — Initial Access: Phishing macro (Invoice_Q1_2026.docm) → WS-FINANCE-01 [T1566.001]\n'
            '3. Apr 7, 14:47 — Execution: PowerShell encoded download dari dl.techsupport-cdn.net [T1059.001]\n'
            '4. Apr 7, 14:49 — C2 Established: svchost_update.exe → 45.77.65.211:443 [T1071.001]\n'
            '5. Apr 8, 09:10 — Discovery: whoami, ipconfig, net group, nltest, systeminfo [T1082]\n'
            '6. Apr 8, 09:15 — Credential Access: Mimikatz → LSASS dump [T1003.001]\n'
            '7. Apr 8, 09:20 — Lateral Movement: Pass-the-Hash → DC-01 [T1550.002]\n'
            '8. Apr 8, 09:22 — Persistence: Create svc_backup + add to Domain Admins [T1136.002]\n'
            '9. Apr 8, 09:25 — Lateral Movement: SSH → DB-SERVER-01 [T1021.004]\n'
            '10. Apr 8, 09:30 — Persistence: Scheduled task SystemUpdate [T1053.005]\n'
            '11. Apr 8, 14:00 — Collection: mysqldump → /tmp/bak.tar.gz [T1005]\n'
            '12. Apr 9, 02:30 — Exfiltration: curl upload 250MB → 185.234.216.88 [T1041]\n'
            '13. Apr 9, 02:35 — Defense Evasion: rm /tmp/bak.tar.gz [T1070.004]\n'
            '14. Apr 9, 02:40 — Defense Evasion: wevtutil cl Security [T1070.001]'
        ),
        'query': '(Kombinasi semua query di atas, disusun kronologis)',
        'explanation': 'Full APT kill chain: Recon → Initial Access → Execution → C2 → Discovery → Credential Access → Lateral Movement → Persistence → Collection → Exfiltration → Defense Evasion. Total durasi: ~41 jam.',
    },
]

# ─── STUDENT VERSION (Questions Only) ───────────────────
for q in questions:
    doc.add_heading(f'Question {q["num"]}', level=2)
    p = doc.add_paragraph()
    r = p.add_run(f'{q["level"]}  |  {q["category"]}')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph(q['q'])
    
    hp = doc.add_paragraph()
    hp.add_run('💡 Hint: ').bold = True
    hr = hp.add_run(q['hint'])
    hr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    hr.font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_paragraph('Answer: _______________________________________________')
    doc.add_paragraph()

doc.add_page_break()

# ─── ANSWER KEY ─────────────────────────────────────────
doc.add_heading('ANSWER KEY', level=1)
p = doc.add_paragraph()
p.add_run('⚠️ JANGAN BUKA HALAMAN INI SEBELUM MENCOBA SENDIRI!').bold = True
p.add_run('\nJawaban dan penjelasan di bawah ini adalah untuk instruktur atau setelah kamu selesai mengerjakan.')

doc.add_paragraph()

for q in questions:
    doc.add_heading(f'Q{q["num"]}: {q["category"]}', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Question: ').bold = True
    p.add_run(q['q'])
    
    p2 = doc.add_paragraph()
    p2.add_run('SPL Query:\n').bold = True
    for line in q['query'].split('\n'):
        qp = doc.add_paragraph(line, style='No Spacing')
        qp.runs[0].font.name = 'Consolas'
        qp.runs[0].font.size = Pt(9)
        qp.runs[0].font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)
    
    p3 = doc.add_paragraph()
    p3.add_run('Answer: ').bold = True
    p3.add_run(q['answer'])
    
    p4 = doc.add_paragraph()
    p4.add_run('Explanation: ').bold = True
    p4.add_run(q['explanation'])
    
    doc.add_paragraph()

# ─── SAVE ────────────────────────────────────────────────
output = "/mnt/d/cheatsheet-blue-team/dataset/SOC_Cyber_Drill_Questions.docx"
doc.save(output)
print(f"✅ Questions saved to: {output}")
